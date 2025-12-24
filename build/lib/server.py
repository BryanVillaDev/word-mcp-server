# math_server.py
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.section import Section
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from mcp.server.fastmcp import FastMCP
import cv2
import numpy as np
from io import BytesIO
import os
import json
from typing import Dict, List, Any, Optional
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE

mcp = FastMCP("Word MCP Server", "1.0")

document = Document()

# Khởi tạo cấu trúc dữ liệu cho Resources và Prompts
RESOURCES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resources")
PROMPTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "prompts")

# Tạo thư mục resources và prompts nếu chưa tồn tại
os.makedirs(RESOURCES_DIR, exist_ok=True)
os.makedirs(PROMPTS_DIR, exist_ok=True)

# Dictionary lưu trữ tài nguyên và prompt trong bộ nhớ
resources_cache = {}


@mcp.tool()
def save_file(filename: str):
    """
    Save file to disk
    - filename: path where the file should be saved (including filename)
    """
    try:
        # Check if the filename has a .docx extension
        if not filename.endswith('.docx'):
            filename = f"{filename}.docx"
            
        # If no directory specified, save to resources directory
        if not os.path.dirname(filename):
            filename = os.path.join(RESOURCES_DIR, filename)
            
        # Ensure the directory exists
        directory = os.path.dirname(filename)
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)
            
        # Debug info
        print(f"Attempting to save document to: {filename}")
            
        # Save the document
        document.save(filename)
        
        # Also save to resources for reference
        resource_id = os.path.basename(filename)
        save_resource(resource_id, resource_id)
        
        return f"File saved successfully to: {filename}"
    except Exception as e:
        error_msg = f"Error saving file: {str(e)}"
        print(error_msg)  # Debug print
        return error_msg

@mcp.tool()
def add_heading(content: str, level: int):
    """
    Add heading to the document
        - Content: nội dung title hoặc heading
        - Level: bậc của heading (0, 1, 2, ...). Số càng nhỏ font chữ càng lớn.
    
    """
    document.add_heading(content, level)

@mcp.tool()
def add_paragraph(
    content: str,
    style: str = "Normal",
    font_size: int = 12,
    bold: bool = False,
    italic: bool = False,
    alignment: WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.LEFT,
):
    """
    Add paragraph to the document
        - Content: nội dung paragraph
    """
    p = document.add_paragraph(content)
    p.style = style
    p.alignment = alignment
    run = p.runs[0]
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    return p


@mcp.tool()
def update_paragraph(
    p,
    content: str = None,
    style: str = "Normal",
    font_size: int = 12,
    bold: bool = False,
    italic: bool = False,
    color: str = None,
    alignment: WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.LEFT,
):
    """
    Update paragraph
        - Content: nội dung paragraph
        - Style: style of paragraph
        - Font size: font size of paragraph
        - Bold: bold or not
        - Italic: italic or not
        - Color: màu chữ (black, blue, green, dark blue, dark red, dark yellow, 
          dark green, pink, red, white, teal, yellow, violet, gray25, gray50)
        - Alignment: căn lề (LEFT, RIGHT, CENTER, JUSTIFY)
    """
    p.style = style
    p.alignment = alignment
    
    run = p.runs[0]
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    
    if color is not None:
        from common import color_paragraph
        color_element = color_paragraph(p, color)
        run.font.color.val = color_element

    if content is not None:
        new_run = p.add_run(content)
        new_run.font.size = Pt(font_size)
        new_run.font.bold = bold
        new_run.font.italic = italic
        if color is not None:
            new_run.font.color.val = color_element
            
    return p 


@mcp.tool()
def add_section(section = WD_SECTION_START.NEW_PAGE) -> Section:
    """
    Add section to the document
    """
    return document.add_section(section)

@mcp.tool()
def set_number_of_columns(section, cols):
    """
    Set number of columns for a section
        - Section: section to set columns
        - Cols: number of columns
    """
    section._sectPr.xpath("./w:cols")[0].set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num", str(cols))

@mcp.tool()
def add_run_to_paragraph(
    p,
    content: str,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    color: str = None,
    highlight: str = None
):
    """
    Thêm câu vào đoạn văn đã được khởi tạo
        - p: paragraph đã được khởi tạo
        - content: nội dung cần thêm
        - bold: in đậm hoặc không
        - italic: in nghiêng hoặc không
        - underline: gạch chân hoặc không
        - color: màu chữ (black, blue, green, dark blue, dark red, dark yellow,
          dark green, pink, red, white, teal, yellow, violet, gray25, gray50)
        - highlight: màu nền highlight cho chữ
    """
    sentence_element = p.add_run(str(content))
    sentence_element.bold = bold
    sentence_element.italic = italic
    sentence_element.underline = underline
    
    if color is not None:
        from common import color_paragraph
        color_element = color_paragraph(p, color)
        sentence_element.font.color.val = color_element
    
    if highlight is not None:
        from docx.enum.text import WD_COLOR_INDEX
        
        if highlight == 'black':
            color_element = WD_COLOR_INDEX.BLACK
        elif highlight == 'blue':
            color_element = WD_COLOR_INDEX.BLUE
        elif highlight == 'green':
            color_element = WD_COLOR_INDEX.BRIGHT_GREEN
        elif highlight == 'dark blue':
            color_element = WD_COLOR_INDEX.DARK_BLUE
        elif highlight == 'dark red':
            color_element = WD_COLOR_INDEX.DARK_RED
        elif highlight == 'dark yellow':
            color_element = WD_COLOR_INDEX.DARK_YELLOW
        elif highlight == 'dark green':
            color_element = WD_COLOR_INDEX.GREEN
        elif highlight == 'pink':
            color_element = WD_COLOR_INDEX.PINK
        elif highlight == 'red':
            color_element = WD_COLOR_INDEX.RED
        elif highlight == 'white':
            color_element = WD_COLOR_INDEX.WHITE
        elif highlight == 'teal':
            color_element = WD_COLOR_INDEX.TEAL
        elif highlight == 'yellow':
            color_element = WD_COLOR_INDEX.YELLOW
        elif highlight == 'violet':
            color_element = WD_COLOR_INDEX.VIOLET
        elif highlight == 'gray25':
            color_element = WD_COLOR_INDEX.GRAY_25
        elif highlight == 'gray50':
            color_element = WD_COLOR_INDEX.GRAY_50
        
        style = document.styles.add_style(f"highlight_style_{highlight}", WD_STYLE_TYPE.CHARACTER)
        style.font.highlight_color = color_element
        sentence_element.style = style
    
    return sentence_element

@mcp.tool()
def add_picture(image_path_or_stream, width: float = 5.0):
    """
    Thêm hình ảnh vào tài liệu
        - image_path_or_stream: đường dẫn đến file ảnh hoặc ảnh dạng ma trận
        - width: chiều rộng của ảnh (tính bằng inch)
    """
    if isinstance(image_path_or_stream, str):
        img = cv2.imread(image_path_or_stream)
    else:
        img = np.array(image_path_or_stream)

    is_success, im_buf_arr = cv2.imencode(".jpg", img)
    byte_im = im_buf_arr.tobytes()
    stream = BytesIO(byte_im)
    return document.add_picture(stream, width=Inches(width))

@mcp.tool()
def create_new_document():
    """
    Tạo một tài liệu mới, loại bỏ tài liệu hiện tại
    """
    global document
    document = Document()
    return "Đã tạo tài liệu mới"

@mcp.tool()
def open_document(filepath: str):
    """
    Mở một tài liệu docx có sẵn
        - filepath: đường dẫn đến file docx cần mở
    """
    global document
    try:
        document = Document(filepath)
        return f"Đã mở tài liệu từ {filepath}"
    except Exception as e:
        return f"Lỗi khi mở tài liệu: {str(e)}"

@mcp.tool()
def add_table(rows: int, cols: int, style: str = "Table Grid"):
    """
    Thêm bảng vào tài liệu
        - rows: số hàng
        - cols: số cột
        - style: kiểu bảng
    """
    table = document.add_table(rows=rows, cols=cols)
    table.style = style
    return table

@mcp.tool()
def create_table(rows: int, cols: int, style: str = "Table Grid", headers: List[str] = None):
    """
    Tạo bảng với số hàng và cột chỉ định, có thể thêm tiêu đề
        - rows: số hàng (không bao gồm hàng tiêu đề)
        - cols: số cột
        - style: kiểu bảng ("Table Grid", "Light Grid", "Light Shading", etc.)
        - headers: danh sách các tiêu đề cột (độ dài bằng số cột)
    """
    try:
        # Nếu có headers, thêm 1 hàng cho tiêu đề
        actual_rows = rows
        if headers:
            if len(headers) != cols:
                return f"Lỗi: Số lượng tiêu đề ({len(headers)}) khác với số cột ({cols})"
            actual_rows = rows + 1
            
        # Tạo bảng
        table = document.add_table(rows=actual_rows, cols=cols)
        table.style = style
        
        # Thêm tiêu đề nếu có
        if headers:
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                # Định dạng tiêu đề (in đậm, căn giữa)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        # Trả về thông tin của bảng để có thể sử dụng sau này
        table_info = {
            "table_object": table,
            "rows": actual_rows,
            "cols": cols,
            "has_headers": bool(headers)
        }
        
        # Lưu thông tin bảng vào resources để có thể truy cập sau này
        table_id = f"table_{id(table)}"
        save_resource(table_id, table_info)
        
        return {
            "table_id": table_id,
            "table_object": table,
            "message": f"Đã tạo bảng với {actual_rows} hàng và {cols} cột"
        }
    except Exception as e:
        error_msg = f"Lỗi khi tạo bảng: {str(e)}"
        print(error_msg)
        return error_msg

@mcp.tool()
def update_cell(table, row: int, col: int, content: str):
    """
    Cập nhật nội dung cho một ô trong bảng
        - table: bảng cần cập nhật (đối tượng Table hoặc mô tả chuỗi)
        - row: chỉ số hàng
        - col: chỉ số cột
        - content: nội dung cần cập nhật
    """
    try:
        # Print debug info
        print(f"Updating cell: row={row}, col={col}, table type={type(table)}")
        
        # Xử lý trường hợp table là chuỗi
        if isinstance(table, str):
            print(f"Table is a string: {table}")
            # Sử dụng bảng cuối cùng được thêm vào document
            if not document.tables:
                return "Không tìm thấy bảng nào trong tài liệu"
                
            print(f"Using last table in document. Total tables: {len(document.tables)}")
            real_table = document.tables[-1]
        else:
            real_table = table
        
        # Truy cập vào ô theo cách được khuyến nghị trong tài liệu
        try:
            if row >= len(real_table.rows):
                return f"Lỗi: Chỉ số hàng {row} vượt quá số hàng trong bảng ({len(real_table.rows)})"
                
            row_cells = real_table.rows[row].cells
            
            if col >= len(row_cells):
                return f"Lỗi: Chỉ số cột {col} vượt quá số cột trong hàng ({len(row_cells)})"
                
            cell = row_cells[col]
            
            # Xóa nội dung hiện tại
            cell.text = ""
            
            # Thêm nội dung mới
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(content)
            run.font.size = Pt(10)  # Kích thước font mặc định
            
            return cell
        except Exception as cell_error:
            error_msg = f"Lỗi khi truy cập ô: {str(cell_error)}"
            print(error_msg)
            return error_msg
    except Exception as e:
        error_msg = f"Lỗi khi cập nhật ô: {str(e)}"
        print(error_msg)
        import traceback
        print(traceback.format_exc())
        return error_msg

@mcp.tool()
def add_page_break():
    """
    Thêm ngắt trang
    """
    document.add_page_break()
    return "Đã thêm ngắt trang"

@mcp.tool()
def fill_table_cell(table, row: int, col: int, content: str, bold: bool = False, alignment = None, font_size: int = None):
    """
    Điền nội dung vào ô trong bảng với định dạng
        - table: đối tượng bảng hoặc table_id (chuỗi)
        - row: chỉ số hàng
        - col: chỉ số cột
        - content: nội dung cần thêm
        - bold: in đậm hay không
        - alignment: căn lề (LEFT, RIGHT, CENTER)
        - font_size: kích thước font
    """
    try:
        # Print debug info
        print(f"Filling cell: row={row}, col={col}, content={content}, table type={type(table)}")
        
        # Xử lý trường hợp table là chuỗi (table_id)
        if isinstance(table, str):
            print(f"Table is a string ID: {table}")
            
            # Trường hợp 1: Đây là table_id từ hàm create_table
            if table.startswith('table_'):
                table_info = get_resource(table)
                if isinstance(table_info, dict) and "table_object" in table_info:
                    real_table = table_info["table_object"]
                else:
                    # Trường hợp 2: Sử dụng bảng cuối cùng trong tài liệu
                    if not document.tables:
                        return "Không tìm thấy bảng nào trong tài liệu"
                    real_table = document.tables[-1]
            else:
                # Trường hợp 3: Sử dụng bảng cuối cùng trong tài liệu
                if not document.tables:
                    return "Không tìm thấy bảng nào trong tài liệu"
                real_table = document.tables[-1]
        else:
            real_table = table
        
        # Kiểm tra có đủ hàng và cột không
        if row >= len(real_table.rows):
            return f"Lỗi: Chỉ số hàng {row} vượt quá số hàng trong bảng ({len(real_table.rows)})"
            
        if col >= len(real_table.rows[row].cells):
            return f"Lỗi: Chỉ số cột {col} vượt quá số cột trong hàng ({len(real_table.rows[row].cells)})"
            
        # Truy cập ô
        cell = real_table.rows[row].cells[col]
        
        # Clear existing content
        for paragraph in cell.paragraphs:
            paragraph.clear()
        
        # Add new content
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = paragraph.add_run(content)
        
        # Set formatting
        if bold:
            run.bold = True
        
        if font_size:
            run.font.size = Pt(font_size)
        
        if alignment:
            if alignment == "LEFT":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif alignment == "RIGHT":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif alignment == "CENTER":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == "JUSTIFY":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        return paragraph
    except Exception as e:
        error_msg = f"Lỗi khi điền nội dung vào ô: {str(e)}"
        print(error_msg)
        import traceback
        print(traceback.format_exc())
        return error_msg

@mcp.tool()
def add_table_row(table, data: List[str], is_header: bool = False):
    """
    Thêm một hàng vào bảng với dữ liệu được cung cấp
        - table: đối tượng bảng
        - data: danh sách dữ liệu cho từng ô
        - is_header: có phải hàng tiêu đề không
    """
    try:
        # Print debug info
        print(f"Adding row to table: data={data}, is_header={is_header}")
        
        # Add a new row
        row = table.add_row()
        
        # Fill the cells
        for i, content in enumerate(data):
            if i < len(table.columns):
                cell = row.cells[i]
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                run = paragraph.add_run(content)
                
                if is_header:
                    run.bold = True
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        return row
    except Exception as e:
        error_msg = f"Lỗi khi thêm hàng vào bảng: {str(e)}"
        print(error_msg)
        return error_msg

@mcp.tool()
def create_simple_table_with_data(headers: List[str], data: List[List[str]], style: str = "Table Grid"):
    """
    Tạo một bảng hoàn chỉnh với dữ liệu
        - headers: danh sách tiêu đề cột
        - data: danh sách các hàng dữ liệu
        - style: kiểu bảng
    """
    try:
        if not headers or not data:
            return "Lỗi: headers hoặc data không được để trống"
            
        cols = len(headers)
        
        # Kiểm tra xem các hàng dữ liệu có đúng số cột không
        for i, row in enumerate(data):
            if len(row) != cols:
                return f"Lỗi: Hàng {i} có {len(row)} cột nhưng cần {cols} cột"
        
        # Tạo bảng với số hàng bằng số lượng dữ liệu + 1 hàng tiêu đề
        table = document.add_table(rows=1, cols=cols)  # Start with 1 row for headers
        table.style = style
        
        # Điền tiêu đề
        for i, header in enumerate(headers):
            fill_table_cell(table, 0, i, header, bold=True, alignment="CENTER")
        
        # Điền dữ liệu
        for row_data in data:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                fill_table_cell(table, len(table.rows)-1, i, cell_data)
        
        return table
    except Exception as e:
        error_msg = f"Lỗi khi tạo bảng với dữ liệu: {str(e)}"
        print(error_msg)
        return error_msg

# PHẦN RESOURCES - quản lý tài nguyên
@mcp.tool()
def save_resource(resource_id: str, content: Any) -> str:
    """
    Lưu tài nguyên vào bộ nhớ và hệ thống file
    
    - resource_id: định danh duy nhất cho tài nguyên
    - content: nội dung của tài nguyên (chuỗi văn bản, đường dẫn đến file, hoặc dữ liệu JSON)
    
    Trả về: Thông báo kết quả
    """
    try:
        # Print debug info
        print(f"Saving resource: '{resource_id}', content type: {type(content)}")
        
        # Lưu vào cache bộ nhớ
        resources_cache[resource_id] = content
        
        # Lưu vào file
        resource_path = os.path.join(RESOURCES_DIR, f"{resource_id}.json")
        
        # If content is a file path and exists
        if isinstance(content, str) and os.path.exists(content) and not os.path.isdir(content):
            try:
                # If it's a docx file, we'll just store the reference
                if content.endswith('.docx'):
                    with open(resource_path, 'w', encoding='utf-8') as f:
                        json.dump({"content": content, "type": "docx_file"}, f, ensure_ascii=False, indent=2)
                else:
                    # For other file types, we might want to store the content directly
                    with open(resource_path, 'w', encoding='utf-8') as f:
                        json.dump({"content": content, "type": "file_path"}, f, ensure_ascii=False, indent=2)
            except Exception as file_error:
                print(f"Error handling file content: {file_error}")
                raise
        else:
            # Regular content (string, dict, etc.)
            with open(resource_path, 'w', encoding='utf-8') as f:
                if isinstance(content, (dict, list)):
                    json.dump(content, f, ensure_ascii=False, indent=2)
                else:
                    json.dump({"content": str(content)}, f, ensure_ascii=False, indent=2)
        
        return f"Đã lưu tài nguyên '{resource_id}' thành công"
    except Exception as e:
        error_msg = f"Lỗi khi lưu tài nguyên: {str(e)}"
        print(error_msg)  # Debug print
        return error_msg

@mcp.tool()
def get_resource(resource_id: str) -> Any:
    """
    Lấy tài nguyên từ bộ nhớ cache hoặc file hệ thống
    
    - resource_id: định danh của tài nguyên cần lấy
    
    Trả về: Nội dung của tài nguyên hoặc thông báo lỗi
    """
    # Kiểm tra trong cache
    if resource_id in resources_cache:
        return resources_cache[resource_id]
    
    # Không có trong cache, thử đọc từ file
    resource_path = os.path.join(RESOURCES_DIR, f"{resource_id}.json")
    try:
        if os.path.exists(resource_path):
            with open(resource_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # Nếu dữ liệu được lưu dưới dạng đơn giản, lấy trường content
                if isinstance(data, dict) and len(data) == 1 and "content" in data:
                    content = data["content"]
                else:
                    content = data
                
                # Cập nhật vào cache
                resources_cache[resource_id] = content
                return content
        else:
            return f"Không tìm thấy tài nguyên '{resource_id}'"
    except Exception as e:
        return f"Lỗi khi đọc tài nguyên: {str(e)}"

@mcp.tool()
def list_resources() -> List[str]:
    """
    Liệt kê danh sách tất cả tài nguyên có sẵn
    
    Trả về: Danh sách các định danh tài nguyên
    """
    # Lấy danh sách từ thư mục
    resources = []
    if os.path.exists(RESOURCES_DIR):
        for filename in os.listdir(RESOURCES_DIR):
            if filename.endswith('.json'):
                resources.append(filename[:-5])  # Bỏ phần mở rộng .json
    
    return resources

@mcp.tool()
def delete_resource(resource_id: str) -> str:
    """
    Xóa tài nguyên
    
    - resource_id: định danh của tài nguyên cần xóa
    
    Trả về: Thông báo kết quả
    """
    # Xóa khỏi cache
    if resource_id in resources_cache:
        del resources_cache[resource_id]
    
    # Xóa file
    resource_path = os.path.join(RESOURCES_DIR, f"{resource_id}.json")
    try:
        if os.path.exists(resource_path):
            os.remove(resource_path)
            return f"Đã xóa tài nguyên '{resource_id}' thành công"
        else:
            return f"Không tìm thấy tài nguyên '{resource_id}'"
    except Exception as e:
        return f"Lỗi khi xóa tài nguyên: {str(e)}"

# PHẦN PROMPT - quản lý templates và prompts
@mcp.tool()
def save_prompt(prompt_id: str, template: str, description: str = "", metadata: Dict = None) -> str:
    """
    Lưu prompt template
    
    - prompt_id: định danh duy nhất cho prompt
    - template: nội dung mẫu của prompt, có thể chứa biến dạng {variable_name}
    - description: mô tả về mục đích và cách sử dụng của prompt
    - metadata: thông tin bổ sung về prompt (tags, tác giả, v.v.)
    
    Trả về: Thông báo kết quả
    """
    prompt_data = {
        "template": template,
        "description": description,
        "metadata": metadata or {},
        "created_at": str(import_datetime_and_get_now())
    }
    
    # Lưu vào file
    prompt_path = os.path.join(PROMPTS_DIR, f"{prompt_id}.json")
    try:
        with open(prompt_path, 'w', encoding='utf-8') as f:
            json.dump(prompt_data, f, ensure_ascii=False, indent=2)
        return f"Đã lưu prompt '{prompt_id}' thành công"
    except Exception as e:
        return f"Lỗi khi lưu prompt: {str(e)}"

@mcp.tool()
def get_prompt(prompt_id: str) -> Dict:
    """
    Lấy thông tin về một prompt
    
    - prompt_id: định danh của prompt cần lấy
    
    Trả về: Thông tin đầy đủ về prompt hoặc thông báo lỗi
    """
    prompt_path = os.path.join(PROMPTS_DIR, f"{prompt_id}.json")
    try:
        if os.path.exists(prompt_path):
            with open(prompt_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        else:
            return {"error": f"Không tìm thấy prompt '{prompt_id}'"}
    except Exception as e:
        return {"error": f"Lỗi khi đọc prompt: {str(e)}"}

@mcp.tool()
def list_prompts() -> List[str]:
    """
    Liệt kê danh sách tất cả prompts có sẵn
    
    Trả về: Danh sách các định danh prompt
    """
    prompts = []
    if os.path.exists(PROMPTS_DIR):
        for filename in os.listdir(PROMPTS_DIR):
            if filename.endswith('.json'):
                prompts.append(filename[:-5])  # Bỏ phần mở rộng .json
    
    return prompts

@mcp.tool()
def delete_prompt(prompt_id: str) -> str:
    """
    Xóa prompt
    
    - prompt_id: định danh của prompt cần xóa
    
    Trả về: Thông báo kết quả
    """
    prompt_path = os.path.join(PROMPTS_DIR, f"{prompt_id}.json")
    try:
        if os.path.exists(prompt_path):
            os.remove(prompt_path)
            return f"Đã xóa prompt '{prompt_id}' thành công"
        else:
            return f"Không tìm thấy prompt '{prompt_id}'"
    except Exception as e:
        return f"Lỗi khi xóa prompt: {str(e)}"

@mcp.tool()
def render_prompt(prompt_id: str, variables: Dict = None) -> str:
    """
    Render một prompt với các biến được cung cấp
    
    - prompt_id: định danh của prompt
    - variables: từ điển chứa các giá trị thay thế cho các biến trong template
    
    Trả về: Prompt đã được render với các biến được thay thế
    """
    prompt_info = get_prompt(prompt_id)
    
    if "error" in prompt_info:
        return prompt_info["error"]
    
    template = prompt_info.get("template", "")
    
    # Thực hiện thay thế biến
    if variables:
        try:
            # Sử dụng format string của Python
            for key, value in variables.items():
                placeholder = "{" + key + "}"
                template = template.replace(placeholder, str(value))
        except Exception as e:
            return f"Lỗi khi render prompt: {str(e)}"
    
    return template

def import_datetime_and_get_now():
    """Helper function to get current datetime"""
    from datetime import datetime
    return datetime.now()

if __name__ == "__main__":
    print("Server đang khởi động...")
    mcp.run(transport="stdio")